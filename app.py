import streamlit as st
import zipfile
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
import pandas as pd
import os

def create_word_document(image_data, df, index_col, image_col, max_width, max_height, output_file):
    # Create a new Word document
    doc = Document()

    # Check if specified columns exist
    if index_col not in df.columns or image_col not in df.columns:
        st.error(f"Excel file must contain '{index_col}' and '{image_col}' columns.")
        return

    # Loop through the rows in the DataFrame and add images to the document
    for _, row in df.iterrows():
        index = row[index_col]
        image_base_name = row[image_col]

        # Find the corresponding image data from the dictionary
        img_data = None
        for filename, data in image_data.items():
            if os.path.splitext(filename)[0] == image_base_name:
                img_data = data
                break

        if img_data is None:
            st.error(f"Image file {image_base_name} not found in the ZIP archive.")
            continue

        # Open the image to get its size
        with Image.open(io.BytesIO(img_data)) as img:
            # Calculate the resize ratio while maintaining the aspect ratio
            ratio = min(max_width / img.width, max_height / img.height)
            width = img.width * ratio
            height = img.height * ratio

        # Add the index and file name
        heading = doc.add_heading(f'{image_base_name}', level=1)
        heading.alignment = 1  # Center align

        # Add the image
        image_stream = io.BytesIO(img_data)
        doc.add_picture(image_stream, width=Inches(width))

        # Center the picture
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = 1  # Center align

        # Add a paragraph break
        doc.add_paragraph('\n')

        # Add a page break after each image except the last one
        if index < df[index_col].max():
            doc.add_page_break()

    # Save the document
    doc.save(output_file)
    return output_file

# --- Streamlit App Code ---
st.title("Image to Word Document Converter")

# Input fields for user to specify the maximum dimensions
max_width = st.number_input("Maximum Width (in inches)", value=5.0)
max_height = st.number_input("Maximum Height (in inches)", value=5.0)

# File uploaders
uploaded_zip = st.file_uploader("Upload a ZIP file containing images", type="zip")
uploaded_excel = st.file_uploader("Upload an Excel file with image indexes and names", type="xlsx")

# Option to specify the index and image columns after uploading Excel file
index_col, image_col = None, None
if uploaded_excel:
    try:
        df = pd.read_excel(uploaded_excel)
        columns = df.columns.tolist()
        index_col = st.selectbox("Select the column for index", options=columns)
        image_col = st.selectbox("Select the column for image file names", options=columns)
    except Exception as e:
        st.error(f"Error reading Excel file: {e}")

# Button to generate the document
if st.button("Generate Document"):
    if uploaded_zip is None:
        st.error("Please upload a ZIP file containing images.")
    elif uploaded_excel is None:
        st.error("Please upload an Excel file with image indexes and names.")
    elif index_col is None or image_col is None:
        st.error("Please select the appropriate columns from the Excel file.")
    else:
        # Extract ZIP file to an in-memory buffer
        with zipfile.ZipFile(uploaded_zip, 'r') as zip_ref:
            image_data = {file_info.filename: zip_ref.read(file_info) for file_info in zip_ref.infolist() if not file_info.is_dir()}

        # Generate the Word document
        output_file = "images_document_final.docx"
        create_word_document(image_data, df, index_col, image_col, max_width, max_height, output_file)
        st.success(f"Document created successfully: {output_file}")
        st.download_button("Download Document", data=open(output_file, "rb"), file_name=output_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.write("Upload a ZIP file containing the images and an Excel file with the image indexes and names.")
