import streamlit as st
import os
import zipfile
from docx import Document
from docx.shared import Inches
from PIL import Image
import tempfile

def create_word_document(image_folder, max_width, max_height, output_file):
    # Create a new Word document
    doc = Document()

    # List all files in the image folder
    image_files = sorted([f for f in os.listdir(image_folder) if os.path.isfile(os.path.join(image_folder, f))])

    # Loop through the image files and add them to the document with proper indexing
    for index, image_file in enumerate(image_files, start=1):
        # Open the image to get its size
        img_path = os.path.join(image_folder, image_file)
        with Image.open(img_path) as img:
            # Calculate the resize ratio while maintaining the aspect ratio
            ratio = min(max_width / img.width, max_height / img.height)
            width = img.width * ratio
            height = img.height * ratio

        # Add the index and file name
        heading = doc.add_heading(f'{index}. {image_file}', level=1)
        heading.alignment = 1  # Center align

        # Add the image
        doc.add_picture(img_path, width=Inches(width))

        # Center the picture
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = 1  # Center align

        # Add a paragraph break
        doc.add_paragraph('\n')

        # Add a page break after each image except the last one
        if index < len(image_files):
            doc.add_page_break()

    # Save the document
    doc.save(output_file)
    return output_file

st.title("Image to Word Document Converter")

# Input fields for user to specify the maximum dimensions
max_width = st.number_input("Maximum Width (in inches)", value=5.0)
max_height = st.number_input("Maximum Height (in inches)", value=5.0)

# File uploader to upload a ZIP file
uploaded_file = st.file_uploader("Upload a ZIP file containing images", type="zip")

# Button to generate the document
if st.button("Generate Document"):
    if uploaded_file is None:
        st.error("Please upload a ZIP file containing images.")
    else:
        # Create a temporary directory to extract the ZIP file
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Extract the ZIP file
            with zipfile.ZipFile(uploaded_file, 'r') as zip_ref:
                zip_ref.extractall(tmp_dir)

            # Generate the Word document
            output_file = "images_document_final.docx"
            create_word_document(tmp_dir, max_width, max_height, output_file)
            st.success(f"Document created successfully: {output_file}")
            st.download_button("Download Document", data=open(output_file, "rb"), file_name=output_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.write("Upload a ZIP file containing the images you want to include in the Word document.")
